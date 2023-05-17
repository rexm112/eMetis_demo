import bs4
import requests
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer
from pdfminer.high_level import extract_text
from docx import Document

# Simple version
def pdf_to_text_simple(path):
    text = extract_text(path)
    return text

# 解析Word文档并返回text
def docx_to_text_simple(path):
    doc = Document(path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text


# 解析Word文档带段落，目前未使用
def docx_to_text(path):
    doc = Document(path)
    paragraphs_content = {}
    for i, para in enumerate(doc.paragraphs, start=1):
        if para.text.strip():  # 只有在段落有内容的情况下才添加
            paragraphs_content[i] = para.text
    return paragraphs_content


# 解析PDF文档
def pdf_to_text(path):
    pages_content = {}
    for page_num, page_layout in enumerate(extract_pages(path), start=1):
        text = ""
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                text += element.get_text()
        if text.strip():  # only add page if there's content
            pages_content[page_num] = text
    return pages_content


# 解析URL
def parse_url(url):
    response = requests.get(url)
    soup = bs4.BeautifulSoup(response.content.decode('utf-8'), 'html.parser')
    title = soup.title.string
    paragraphs = soup.find_all('p')
    paragraph_list = []
    paragraph_texts = ""
    for p in paragraphs:
        paragraph_list.append(p.get_text())
    paragraph_texts = paragraph_texts.join(paragraph_list)
    return title, paragraph_texts


docx_path = "./assets/半佛.docx"  # 替换为你的 Word 文件路径
paragraphs_content = docx_to_text(docx_path)
text = docx_to_text_simple(docx_path)

print(text)

""" for para_num, content in paragraphs_content.items():
    print(f"段落 {para_num}:\n{content}\n{'-'*40}\n") """


#url = "https://news.sina.com.cn/c/xl/2023-05-17/doc-imyuauqy5897453.shtml"
#title, text = parse_url(url)
